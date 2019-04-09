#Document Enhancer

import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import random
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

import os
import sys
import torch
import random
import argparse
import numpy as np
from gpt2Pytorch.GPT2.model import (GPT2LMHeadModel)
from gpt2Pytorch.GPT2.utils import load_weight
from gpt2Pytorch.GPT2.config import GPT2Config
from gpt2Pytorch.GPT2.sample import sample_sequence
from gpt2Pytorch.GPT2.encoder import get_encoder

from gpt2Pytorch.mainLib import *


documentold = Document()
paragraph = 0
word = 0


document = Document("test.docx")
document = Document()






def formatStyles():
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0,0,0)
    font.underline = False
    style.paragraph_format.line_spacing = 2
    """
    styles = document.styles
    new_heading_style = styles.add_style('New Heading', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = document.styles['Heading 1']
    font = style.font

    #style = document.styles['Heading 1']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(255,0,0)
    font.underline = False
    """

    styles = document.styles
    styles['Title'].delete()
    style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)

    style = document.styles['Title']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(20)
    font.color.rgb = RGBColor(0,0,0)
    font.underline = False
    style.paragraph_format.line_spacing = 2

def addheader():
    section = document.sections[0]
    header = section.header
    head = header.paragraphs[0]
    list = readParagraph(0)
    #document.add_paragraph("Testing", )
    #name = list[0]
    head.text = list.split( )[1]
    #head.text = readWord(0,1)
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def addheading():
    list = readHeading()
    #headingDate()
    #for i in list:
    #    heading = document.add_paragraph(i)
    heading = document.add_paragraph(list[0])
    heading = document.add_paragraph(list[1])
    heading = document.add_paragraph(headingDate(list[2]))
    #print list[2]
    heading = document.add_paragraph(list[3])

    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.line_spacing = 2


def addtitle():
    #document.add_heading('Document Title', 0)
    title = document.add_paragraph(readParagraph(4))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

def addbody():

    list = readBody()
    for i in list:
        #print(list[i])
        paragraph = document.add_paragraph('\t' + i)
    for paragraph_text in AIconverter(readParagraph(5)).split('\n\n'):
        #print(paragraph_text.strip())
        paragraph = document.add_paragraph("\t"+paragraph_text.strip())
        #body.append(paragraph(paragraph_text.strip()))

    #paragraph = document.add_paragraph(AItext)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 2


def readFile():
    file = open("test.txt", "r")
    content = file.read()

def readHeading():
    doc = docx.Document('test.docx')
    headingList = [doc.paragraphs[0].text, doc.paragraphs[1].text, doc.paragraphs[2].text, doc.paragraphs[3].text]
    return headingList;

def readParagraph(paragraph):
    doc = docx.Document('test.docx')
    return doc.paragraphs[paragraph].text


def headingDate(date):
    #date = readParagraph(2)
    #print(date)
    #date="mar 2020 24"
    #x = re.search("[0-2][0-9]|[1-9]" , date)
    try:
        day = re.search("([^\d])([0-2]|[0-2][0-9])([^\d])" , " "+date+" ")
        print(day.group())
        year = re.search("[2-9][0-9][0-9][0-9]" , date)
        print(year.group())
        month = re.search("[^\s\d][^\s\d][^\s\d]" , date)
        print(month.group())
        newdate=day.group()[1:-1]+" "+month.group().capitalize()+". "+year.group()
        print(newdate)
    except:
        print("error")
        newdate=date
    return newdate

def readBody():
    doc = docx.Document('test.docx')
    list = []
    i=len(doc.paragraphs)-1
    while i < len(doc.paragraphs):
        list.append(doc.paragraphs[i].text)
        i+=1
    return list

#returns true if a word is ignorable and false if important
def ignorable(word):
    ignore = ["The", "the", "To", "to", "Of", "of", "Be", "be", "and", "A", "a", "That", "that", "Have", "have", "I",
              "It", "it", "For", "for", "Not", "not", "With", "with", "You", "you", "As", "as", "Do", "do", "At", "at"
              "This", "this", "By", "by", "or", "An", "an", "From", "from", "Will", "will", "Is", "is"]
    for x in range( 0, len(ignore) ):
        if(word == ignore[x]):
            return True
    return False

#takes in a string and returns a dictionary on the word count of each word
def getRepetitive( text ):
    unique = {}

    for word in text:
        if ignorable(word) == False:
            if len(unique) == 0:
                unique[word] = 1
            else:
                for y in list(unique):
                    if(word == y):
                        unique[word] += 1
                        break
                    unique[word] = 1
    return unique




#MAIN
def main():
    formatStyles()
    addheader()
    addheading()
    addtitle()
    addbody()
    #AIconverter(readParagraph(5))
    document.save('main2.docx')

main()
