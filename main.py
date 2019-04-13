from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import docx
import random
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

document = Document()


def formatStyles():
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)


def addHeader():
    section = document.sections[0]
    header = section.header
    head = header.paragraphs[0]

    list = readHeading()
    name = list[0]
    head.text = name.split()[1]
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def addHeading():
    list = readHeading()
    for i in list:
        heading = document.add_paragraph(i)

    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.line_spacing = 2


def addBody():
    title = document.add_paragraph(readTitle())
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    list = readBody()
    for i in list:
        paragraph = document.add_paragraph('\t' + i)

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 2


def readHeading():
    doc = docx.Document('test.docx')
    headingList = [doc.paragraphs[0].text, doc.paragraphs[1].text, doc.paragraphs[2].text, doc.paragraphs[3].text]
    return headingList;


def readTitle():
    doc = docx.Document('test.docx')
    return doc.paragraphs[4].text


def readBody():
    doc = docx.Document('test.docx')
    list = []
    i=len(doc.paragraphs)-1;
    while i < len(doc.paragraphs):
        list.append(doc.paragraphs[i].text)
        i+=1
    return list





formatStyles()
addHeader()
addHeading()
addBody()
document.save('main.docx')





def replaceStuff():
    conciseArr = ["like", "like", "in fact", "always", "now", "now", "currently", "currently", "because", "becasue", "because", "because", "by", "point out", "for", "be able to", "can", "to", "on", "on", "on", "about", "about", "about", "although", "though", "if", "if", "finally", "decide on", "when", "twice", "most", "until", "of", "gone"]

    inflatedArr = ["along the lines of", "in the nature of", "as a matter of fact", "at all times", "at the present time", "at this point in time", "at the present time", "at this point in time", "because of the fact that", "due to the fact that", "for the reason that", "in light of the fact that", "by the means of", "draw your attention to", "for the purpose of", "have the ability", "have the ability to", "in order to", "in regards to", "in reference to", "with reference to", "in regards to", "in reference to", "with reference to", "in spite of the fact that", "in spite of the fact that", "in the event that", "in the situation that", "in the final analysis", "make decisions about", "on the occasion of", "on two seperate occasions", "the majority of", "until such time as", "with reference to", "no longer present"]

    bodyParagraphs = readBody()
    print (bodyParagraphs)

    for i in range(0, len(bodyParagraphs)):
        bodyParagraphs[i] = bodyParagraphs[i].replace('Date', 'replacement')

    print(bodyParagraphs)
